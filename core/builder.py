"""
core/builder.py
===============
Buduje plik wyjściowy DOCX lub TXT z listy przetłumaczonych DocElement.

Zachowuje kolejność elementów i strukturę dokumentu:
  - DOCX: tabele jako natywne tabele Worda, obrazy osadzone, nagłówki H1-H3
  - TXT:  tabele jako ASCII, separatory stron, bez obrazów
"""

import io
from pathlib import Path
from typing import List

from .extractor import DocElement


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def zbuduj_docx(elementy: List[DocElement], sciezka_wyjscia: str) -> None:
    """Buduje plik DOCX zachowując strukturę dokumentu."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn

    doc = Document()

    # Marginesy strony
    for sekcja in doc.sections:
        sekcja.top_margin    = Inches(1.0)
        sekcja.bottom_margin = Inches(1.0)
        sekcja.left_margin   = Inches(1.2)
        sekcja.right_margin  = Inches(1.2)

    for elem in elementy:

        # --- Separator strony ---
        if elem.kind == "page_sep":
            p = doc.add_paragraph(f"{'─' * 18}  Strona {elem.page}  {'─' * 18}")
            run = p.runs[0]
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x9C, 0xA0, 0xB0)
            continue

        # --- Nagłówek ---
        if elem.kind == "heading":
            tekst = elem.translated_text or elem.text
            if tekst.strip():
                doc.add_heading(tekst.strip(), level=elem.level)
            continue

        # --- Akapit lub element listy ---
        if elem.kind in ("paragraph", "list_item"):
            tekst = elem.translated_text or elem.text
            if not tekst.strip():
                continue
            if elem.kind == "list_item":
                doc.add_paragraph(tekst.strip(), style="List Bullet")
            else:
                doc.add_paragraph(tekst.strip())
            continue

        # --- Tabela ---
        if elem.kind == "table":
            komorki = elem.translated_cells or elem.table_data
            if not komorki:
                continue

            wiersze = len(komorki)
            kolumny = max((len(w) for w in komorki), default=1)
            if wiersze == 0 or kolumny == 0:
                continue

            tabela = doc.add_table(rows=wiersze, cols=kolumny)
            try:
                tabela.style = "Table Grid"
            except Exception:
                pass

            for r, wiersz in enumerate(komorki):
                for c, komorka in enumerate(wiersz):
                    if c < kolumny:
                        tabela.cell(r, c).text = str(komorka or "")

            doc.add_paragraph()  # pustka po tabeli
            continue

        # --- Obraz ---
        if elem.kind == "image" and elem.image_data:
            try:
                strumien = io.BytesIO(elem.image_data)
                doc.add_picture(strumien, width=Inches(5.5))
                doc.add_paragraph()
            except Exception:
                pass  # pomiń uszkodzone/nieobsługiwane obrazy

    doc.save(sciezka_wyjscia)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def zbuduj_txt(elementy: List[DocElement], sciezka_wyjscia: str) -> None:
    """Buduje plik TXT z przetłumaczoną treścią."""
    linie: List[str] = []

    for elem in elementy:

        if elem.kind == "page_sep":
            linie.append("")
            linie.append("=" * 60)
            linie.append(f"  STRONA {elem.page}")
            linie.append("=" * 60)
            linie.append("")
            continue

        if elem.kind == "heading":
            tekst = elem.translated_text or elem.text
            if tekst.strip():
                prefiks = "#" * elem.level + " "
                linie.append("")
                linie.append(prefiks + tekst.strip())
                linie.append("")
            continue

        if elem.kind in ("paragraph", "list_item"):
            tekst = elem.translated_text or elem.text
            if tekst.strip():
                linie.append(tekst.strip())
                linie.append("")
            continue

        if elem.kind == "table":
            komorki = elem.translated_cells or elem.table_data
            if komorki:
                linie.append("[TABELA]")
                for wiersz in komorki:
                    linie.append("  " + " | ".join(str(k or "") for k in wiersz))
                linie.append("")
            continue

        if elem.kind == "image":
            linie.append("[OBRAZ – pominięty w formacie TXT]")
            linie.append("")

    with open(sciezka_wyjscia, "w", encoding="utf-8") as f:
        f.write("\n".join(linie))
